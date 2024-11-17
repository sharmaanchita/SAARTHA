import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import os
import joblib
import matplotlib
from docx import Document


def set_page_query_param(page_name):
    st.experimental_set_query_params(page=page_name)
    
set_page_query_param("carbon_emission")

def calculate_motorcycle_emissions(distance, num_days):
    CO2_PER_KM = 0.129
    total_distance = distance * num_days
    carbon_emissions = total_distance * CO2_PER_KM
    return carbon_emissions

def calculate_car_emissions(distance, num_days):
    CO2_PER_KM = 0.1808
    AVG_DISTANCE_PER_DAY = 40.0
    total_distance = distance * num_days
    carbon_emissions = (total_distance / AVG_DISTANCE_PER_DAY) * CO2_PER_KM
    return carbon_emissions

def calculate_public_trans_emissions(distance, num_days):
    CO2_PER_KM = 0.09
    total_distance = distance * num_days
    carbon_emissions = total_distance * CO2_PER_KM
    return carbon_emissions


def gradient_text(text, color1, color2):
    return f"""
    <span style="
        background: linear-gradient(to right, {color1}, {color2});
        -webkit-background-clip: text;
        color: transparent;
    ">
        {text}
    </span>
    """

color1 = "#0d3270"
color2 = "#228B22"
text = "Commute Insights"

styled_text = gradient_text(text, color1, color2)

html_code = f"""
    <div style="background-color: #3CB371; border: 2px solid black; padding: 10px; border-radius: 5px;">
        <h3>{styled_text}</h3>
    </div>
"""
st.markdown(html_code, unsafe_allow_html=True)

st.write('Calculate and reduce your carbon footprint for your daily commute using simple machine learning algorithms.')

st.header(':green[Enter your commute details:]')
data = []
transportation_modes = ['Public transportation', 'Car', 'Motorcycle']
total_emissions = 0

distance_unit = st.selectbox('Distance Unit', ['Kilometers', 'Miles'])

for mode in transportation_modes:
    st.subheader(f'{mode} Commute Details:')
    distance_label = f'Distance ({distance_unit.lower()}) for {mode}:'
    distance_conversion = 1.0

    if distance_unit == 'Miles':
        distance_conversion = 0.621371
        distance_label = f'Distance ({distance_unit.lower()}) for {mode}:'

    distance = st.number_input(distance_label, min_value=0.0, max_value=1000.0, step=0.1, key=f'{mode}_distance')
    num_days = st.number_input(f'Number of days per week for {mode}:', min_value=0, max_value=7, step=1, key=f'{mode}_days')
    distance *= distance_conversion
    if distance < 0 or num_days < 0:
        st.error('Invalid input. Distance and number of days should be non-negative.')
        continue

    if mode == 'Car':
        carbon_emissions = calculate_car_emissions(distance, num_days)
    elif mode == 'Motorcycle':
        carbon_emissions = calculate_motorcycle_emissions(distance, num_days)
    elif mode == 'Public transportation':
        carbon_emissions = calculate_public_trans_emissions(distance, num_days)

    total_emissions += carbon_emissions
    data.append({'Transportation Mode': mode, 'CO2 Emissions (kg)': carbon_emissions})

df = pd.DataFrame(data)
if st.button("Calculate: "):
    st.header('Results:')

    if total_emissions > 10:
        st.subheader(f'Net :red[carbon] :red[emissions]: :red[{format(total_emissions, ".2f")} kg of CO2]')
    else:
        st.subheader(f'Net :green[carbon] :green[emissions]: :green[{format(total_emissions, ".2f")} kg of CO2]')

    max_emissions = max(data, key=lambda x: x['CO2 Emissions (kg)'])['CO2 Emissions (kg)']

    fig, ax = plt.subplots()
    ax.bar(df['Transportation Mode'], df['CO2 Emissions (kg)'])
    ax.set_xlabel('Transportation Mode')
    ax.set_ylabel('CO2 Emissions (kg)')
    ax.set_title('CO2 Emissions by Transportation Mode')

    for i, value in enumerate(df['CO2 Emissions (kg)']):
        ax.text(i, value + max_emissions * 0.02, f'{value:.2f}', ha='center')

    ax.set_ylim(0, max_emissions * 1.4)
    st.pyplot(fig)

    st.subheader('\n\nSustainable transportation recommendations:')
    if total_emissions > 0:
        st.write(':red[Consider the following more sustainable transportation options::]')
        if any(mode in ['Car', 'Motorcycle'] for mode in transportation_modes):
            st.write(':red[- Use public transportation or carpooling when possible:]')
            st.write(':red[- Switch to an electric vehicle:]')
        else:
            st.write(':green[- Keep up the good work!:]')
    else:
        st.write(':green[Your carbon emissions are negligible. Keep up the good work!:]')

st.divider()
color1 = "#0d3270"
color2 = "#228B22"
text = "Trans Sustain"

styled_text = gradient_text(text, color1, color2)

html_code = f"""
    <div style="background-color: #3CB371; border: 2px solid black; padding: 10px; border-radius: 5px;">
        <h3>{styled_text}</h3>
    </div>
"""
st.markdown(html_code, unsafe_allow_html=True)

st.write("Calculate and Log your environmental impact while shipping and commuting using Machine learning algorithms: Logistic Regression.")

matplotlib.use('Agg')

st.header(':green[Enter the details:]')

feature_names_best = ['ton', 'unit', 'transportation_mode', 'c_carbon_neutral', 'eco_friendly', 'delivery', 'ascites','varices', 'items_count', 'avg_cost', 'distance', 'approx_weight', 'mode_number','b_carbon_neutral']

distance_dict = {"Miles":1,"Kilometers":2}
feature_dict = {"No":1,"Yes":2}
transportmodel_dict = {"Rail":1,"Ship":2, "Air":3, "Truck":4, "Car":5}

def get_value(val,my_dict):
	for key,value in my_dict.items():
		if val == key:
			return value

def get_fvalue(val):
	feature_dict = {"No":1,"Yes":2}
	for key,value in feature_dict.items():
		if val == key:
			return value

def get_tvalue(val):
	for key,value in transportmodel_dict.items():
		if val == key:
			return value

# Load ML Models
def load_model(model_file):
	loaded_model = joblib.load(open(os.path.join(model_file),"rb"))
	return loaded_model

doc = Document()
doc.add_heading('Your Carbon Emission', level=1)

st.write("Average shipment weights:")
st.write(":green[Freight tons per railcar - 90:] ")
st.write(":green[Freight tons per truck - 16:]")

transportation_mode = st.radio("Mode of transportation", tuple(transportmodel_dict.keys()))
ton = st.number_input("Tons you wish to ship", 1, 10000)
unit = st.radio("Unit of measurement for distance", tuple(distance_dict.keys()))
distance = st.number_input("Average distance which will be travelled", 1, 1000)
mode_number = 1

st.write("Additional Information about your Shipment!")
items_count = st.number_input("Number of items chosen? (If buying retail)", 0, 100000)
avg_cost = st.number_input("Average cost of products", 0, 100000)

approx_weight = st.number_input("Average weight of packing material used in packaging", 0, 100000)

eco_friendly = st.radio("Have you selected eco-friendly packaging?", tuple(feature_dict.keys()))

delivery = st.radio("Will it be same-day delivery?", tuple(feature_dict.keys()))
ascites = st.radio("Is the product refurbished or new?", tuple(feature_dict.keys()))
varices = st.radio("Can the product be recyled?", tuple(feature_dict.keys()))
b_carbon_neutral = st.radio("Is the business carbon-neutral?", tuple(feature_dict.keys()))
c_carbon_neutral = st.radio("Is the shipment company carbon-neutral?", tuple(feature_dict.keys()))

feature_list = [ton, get_value(unit, distance_dict), get_tvalue(transportation_mode), get_fvalue(c_carbon_neutral),
                get_fvalue(eco_friendly), get_fvalue(delivery), get_fvalue(ascites), get_fvalue(varices),
                items_count, avg_cost, distance, approx_weight, int(mode_number), get_fvalue(b_carbon_neutral)]

pretty_result = {"age": ton, "sex": unit, "steroid": transportation_mode, "antivirals": c_carbon_neutral,
                                 "fatigue": eco_friendly, "spiders": delivery, "ascites": ascites, "varices": varices,
                                 "bilirubin": items_count, "alk_phosphate": avg_cost, "sgot": distance,
                                 "albumin": approx_weight, "protime": mode_number, "histolog": b_carbon_neutral}

single_sample = np.array(feature_list).reshape(1, -1)

model_choice = "LR"

if st.button("Predict"):
    if model_choice == "DecisionTree":
        loaded_model = load_model("models/decision_tree.pkl")
        prediction = loaded_model.predict(single_sample)
        pred_prob = loaded_model.predict_proba(single_sample)
    else:
        loaded_model = load_model("models/logistic_regression.pkl")
        prediction = loaded_model.predict(single_sample)
        pred_prob = loaded_model.predict_proba(single_sample)
		
    doc.add_heading('Chosen mode of transporttaion: '+transportation_mode, level=3)
    doc.add_heading('Tons of shipment you wish to ship: '+str(ton), level=3)
    doc.add_heading('Chosen unit of measurement for distance: '+unit, level=3)
    doc.add_heading('Average distance which will be travelled: '+transportation_mode, level=3)
    doc.add_heading('Additional Information about your Shipment: ', level=2)
    doc.add_heading('Number of items chosen? (If buying retail): '+str(items_count), level=3)
    doc.add_heading('Average cost of products: '+str(avg_cost), level=3)
    doc.add_heading('Average weight of packing material used in packaging: '+str(approx_weight), level=3)
    doc.add_heading('Average weight of packing material used in packaging: '+eco_friendly, level=3)
    doc.add_heading('It will be same-day delivery: '+delivery, level=3)
    doc.add_heading('The product refurbished: '+ascites, level=3)
    doc.add_heading('The product can be recyled: '+varices, level=3)
    doc.add_heading('The business is carbon-neutral: '+b_carbon_neutral, level=3)
    doc.add_heading('The shipment company is carbon-neutral: '+c_carbon_neutral, level=3)
    doc.add_heading('Result', level=1)
	
    if unit == "Miles":
        prescriptive_message_temp ="""
	    <div style="overflow-x: auto; padding:10px;border-radius:5px;margin:10px;">
		<h3 style="text-align:justify;color:white;padding:10px">Options to offset your carbon footprint</h3>
		<ul>
		<li style="text-align:justify;color:white;padding:10px">By not selecting same-day delivery you could reduce your CO2e emissions by an estimated 0.8 tons </li>
		<li style="text-align:justify;color:white;padding:10px">Overall Carbon footprint score is equivalent to 1 car off the road for 1 hour</li>
		<li style="text-align:justify;color:white;padding:10px">Plant one tree today!</li>
		<li style="text-align:justify;color:white;padding:10px">Use refurbuished items instead of new</li>
		<ul>
	    </div>
	    """
        st.warning(":red[You have significant carbon footprint!:]")
        doc.add_heading('You have significant carbon footprint!', level=2)
		
        if(transportation_mode == 'Car'):
            st.write(":red[Total estimated CO2e emissions from all selected modes] "+str(pred_prob[0][0]*0.7))
            st.write(":red[Overall Carbon Footprint score] "+str(items_count*1.2))
            doc.add_heading("Total estimated CO2e emissions from all selected modes "+str(pred_prob[0][0]*0.7)+"\nOverall Carbon Footprint score "+str(items_count*1.2), level=3)
            st.subheader("What can you do to offset this carbon footprint?")
            st.markdown(prescriptive_message_temp, unsafe_allow_html=True)
        
        else:
            doc.add_heading("Total estimated CO2e emissions from all selected modes "+str(pred_prob[0][0]*0.2)+"\nOverall Carbon Footprint score "+str(items_count*0.4), level=3)
            st.write(":red[Total estimated CO2e emissions from all selected modes] "+str(pred_prob[0][0]*0.2))
            st.write(":red[Overall Carbon Footprint score] "+str(items_count*0.4))
            st.subheader("What can you do to offset this carbon footprint?")
            st.markdown(prescriptive_message_temp, unsafe_allow_html=True)
            doc.add_heading("Options to offset your carbon footprint",level=2)
            doc.add_heading("1. By not selecting same-day delivery you could reduce your CO2e emissions by an estimated 0.8 tons", level=3)
            doc.add_heading("2. Overall Carbon footprint score is equivalent to 1 car off the road for 1 hour", level=3)
            doc.add_heading("3. Use refurbuished items instead of new", level=3)
            doc.add_heading("4. Plant one tree today!", level=3)
    else:
        st.success("You do not have significant carbon footprint results")
        doc.add_heading("You do not have significant carbon footprint results", level=2)
        doc.add_heading("Performance Score impact "+str(items_count),level=3)
        st.write(":green[Performance Score Impact:] "+str(items_count))
        st.subheader("Prediction Probability Score using {}".format(model_choice))

    doc.save('Report.doc')
    st.download_button(
    label="Download Report",
    data=open("Report.doc", "rb").read(),
    file_name="Report.doc",
    mime="application/octet-stream",
    help="Click to download the carbon emission."
    )
    
st.divider()
color1 = "#0d3270"
color2 = "#228B22"
text = "Carbon Footprint Calculator"

styled_text = gradient_text(text, color1, color2)

html_code = f"""
    <div style="background-color: #3CB371; border: 2px solid black; padding: 10px; border-radius: 5px;">
        <h3>{styled_text}</h3>
    </div>
"""
st.markdown(html_code, unsafe_allow_html=True)  

st.write("Illuminating global carbon footprints through interactive data visualizations.")

df = pd.read_csv("data/carbonfootprint_countries.csv")
st.dataframe(df)

freq_df = pd.read_csv("data/countries_dataset.csv")
st.bar_chart(freq_df['count'])

df['Region'].value_counts().plot(kind='bar')
if st.checkbox("Area Chart"):
    all_columns = df.columns.to_list()
    feat_choices = st.multiselect("Choose a Feature", all_columns)
    new_df = df[feat_choices]
    st.area_chart(new_df)
        
        
        
        

footer="""<style>

a:hover,  a:active {
color: red;
background-color: transparent;
text-decoration: underline;
}

.footer {
position: fixed;
left: 0;
bottom: 0;
width: 100%;
background-color: white;
color: black;
text-align: center;
}
</style>
<div class="footer">
<p>Developed by <a style='display: inline; text-align: center;' href="https://www.linkedin.com/in/harshavardhan-bajoria/" target="_blank">Team SAARTHA</a></p>
</div>
"""
st.markdown(footer,unsafe_allow_html=True)
